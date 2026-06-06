"""
导出功能 —— Word 文档下载
"""
import io
import re
from urllib.parse import quote
from fastapi import APIRouter, Query, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from src.auth_service.dependencies import require_role

router = APIRouter(prefix="/api/export", tags=["文件导出"])


class ExportChatRequest(BaseModel):
    messages: list[dict]


class ExportTextRequest(BaseModel):
    content: str
    title: str = "合同草案"


def _build_docx(title: str, content_blocks: list[tuple[str, str]],
                is_contract: bool = False) -> io.BytesIO:
    """生成格式规范的 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # -- 页面设置 --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # -- 默认样式 --
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # -- 标题 --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'SimHei'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    title_para.paragraph_format.space_after = Pt(20)

    # -- 元信息 --
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"生成时间：{__import__('datetime').datetime.now().strftime('%Y年%m月%d日')}    智能法务系统 生成"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    meta.paragraph_format.space_after = Pt(4)

    disclaimer_para = doc.add_paragraph()
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disclaimer_para.add_run("⚠ 本文档为 AI 辅助生成，不构成正式法律意见，使用前请咨询执业律师审核。")
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    disclaimer_para.paragraph_format.space_after = Pt(16)

    # 分隔线
    sep = doc.add_paragraph()
    sep_run = sep.add_run("—" * 40)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(12)

    for heading, body in content_blocks:
        if heading:
            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.name = 'SimHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(8)

        # 解析正文：识别编号（第X条、X.、一、等）
        paragraphs = _parse_body_paragraphs(body)

        for ptext, is_article, is_bold_key in paragraphs:
            ptext = ptext.strip()
            if not ptext:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(4)

            if is_article:
                # 条款标题 —— 加粗
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(8)
                run = p.add_run(ptext)
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'SimHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            elif is_contract:
                # 合同正文段落 —— 首行缩进2字符
                p.paragraph_format.first_line_indent = Cm(0.74)
                _add_formatted_run(p, ptext)
            else:
                p.paragraph_format.first_line_indent = Cm(0.74)
                _add_formatted_run(p, ptext)

    # -- 文档末尾 --
    doc.add_paragraph()  # 空行
    end_note = doc.add_paragraph()
    end_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_note.add_run("— 文档结束 —")
    end_run.font.size = Pt(9)
    end_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _parse_body_paragraphs(body: str) -> list[tuple[str, bool, bool]]:
    """解析正文段落，识别条款标题和普通段落"""
    # 先清理 markdown 标记
    body = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", body)
    body = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", body)
    body = re.sub(r"#+\s*", "", body)

    lines = body.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # 检测条款标题：第X条、第X章、第X节
        is_article = bool(re.match(r'^第[一二三四五六七八九十百千万\d]+[章节条]', stripped))
        # 检测编号列表
        is_numbered = bool(re.match(r'^[（(]?\d+[)）]', stripped))
        result.append((stripped, is_article, False))
    return result


def _add_formatted_run(para, text: str):
    """添加带格式的文本段落，处理粗体标记"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # 分割粗体部分
    parts = re.split(r'(\*\*(.+?)\*\*)', text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if part.startswith('**') and i + 1 < len(parts):
            # 粗体部分
            bold_text = parts[i + 1]
            run = para.add_run(bold_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            i += 2
        elif part and not part.startswith('**'):
            run = para.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        i += 1


@router.get("/contract-review/{contract_id}")
async def export_contract_review(contract_id: str, current_user: dict = Depends(require_role("admin","legal","business"))):
    """导出合同审查报告为 Word"""
    from src.contract_service.service import contract_service

    contract = contract_service.get_contract(contract_id)
    if not contract:
        return {"error": "合同不存在"}
    if not contract.clauses or all(not cl.risk_level or cl.risk_level.value == "none" for cl in contract.clauses):
        return {"error": "该合同尚未审查或未发现风险，请先进行审查"}

    blocks: list[tuple[str, str]] = [(
        "合同信息",
        f"名称：{contract.title}\n"
        f"类型：{contract.contract_type}\n"
        f"甲方：{contract.party_a}\n乙方：{contract.party_b}",
    )]

    high = sum(1 for cl in (contract.clauses or []) if cl.risk_level and cl.risk_level.value == "high")
    medium = sum(1 for cl in (contract.clauses or []) if cl.risk_level and cl.risk_level.value == "medium")
    low = sum(1 for cl in (contract.clauses or []) if cl.risk_level and cl.risk_level.value == "low")
    blocks.append(("审查摘要", f"共审查 {len(contract.clauses or [])} 条，高风险 {high}，中风险 {medium}，低风险 {low}"))

    for cl in (contract.clauses or []):
        if not cl.risk_level or cl.risk_level.value == "none":
            continue
        parts = []
        if cl.content: parts.append(f"原文：{cl.content[:500]}")
        if cl.risk_analysis: parts.append(f"风险分析：{cl.risk_analysis}")
        if cl.law_basis: parts.append(f"法律依据：{cl.law_basis}")
        if cl.suggestion: parts.append(f"修改建议：{cl.suggestion}")
        blocks.append((f"{cl.clause_number}（{cl.risk_level.value}）", "\n".join(parts)))

    buf = _build_docx(f"合同审查报告 —— {contract.title}", blocks)
    filename = f"审查报告_{contract.title}_{contract_id}.docx"
    safe_name = quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
    )


@router.post("/contract-draft")
async def export_contract_draft(req: ExportTextRequest, current_user: dict = Depends(require_role("admin","legal","business"))):
    """导出 AI 生成的合同草案为 Word（规范格式）"""
    content = req.content.strip()
    if not content:
        return {"error": "合同内容为空"}

    blocks: list[tuple[str, str]] = []
    # 按 ## / ### / 第X条 分块
    sections = re.split(r"\n(?=#{2,3}\s|第[一二三四五六七八九十百千万\d]+条)", content)
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        lines = sec.split("\n", 1)
        first = lines[0].strip()
        if first.startswith("#"):
            heading = re.sub(r"^#+\s*", "", first)
            body = lines[1].strip() if len(lines) > 1 else ""
        elif re.match(r'^第[一二三四五六七八九十百千万\d]+条', first):
            heading = first
            body = lines[1].strip() if len(lines) > 1 else ""
        else:
            heading = ""
            body = sec
        body = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", body)
        body = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", body)
        blocks.append((heading, body))

    if not blocks:
        blocks = [("合同草案", content)]

    try:
        buf = _build_docx(req.title, blocks, is_contract=True)
    except Exception as e:
        return {"error": f"文档生成失败: {str(e)}"}

    safe_name = quote(f"{req.title}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
    )


@router.post("/consultation")
async def export_consultation(req: ExportChatRequest, current_user: dict = Depends(require_role("admin","legal","business"))):
    """导出咨询对话为 Word"""
    blocks: list[tuple[str, str]] = []
    for i, m in enumerate(req.messages):
        role = "用户" if m.get("role") == "user" else "AI 助手"
        blocks.append((f"{role} · 第{i//2+1}轮", m.get("content", "")))

    buf = _build_docx("法律咨询记录", blocks)
    filename = "咨询记录.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})
