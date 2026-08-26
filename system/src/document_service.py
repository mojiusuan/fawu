"""
文书模板与组装服务
"""
import json
import uuid
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from src.config import settings


class DocumentService:

    def __init__(self):
        self._store_path = Path(settings.BASE_DIR) / "data" / "doc_templates.json"
        self._store_path.parent.mkdir(parents=True, exist_ok=True)
        self._export_dir = Path(settings.EXPORT_DIR)
        self._export_dir.mkdir(parents=True, exist_ok=True)
        self._ensure_default_templates()

    def _ensure_default_templates(self):
        if self._store_path.exists():
            return
        templates = [
            {
                "template_id": "civil_complaint_contract",
                "category": "起诉类",
                "title": "民事起诉状（合同纠纷）",
                "description": "适用于买卖合同、服务合同等合同违约纠纷的起诉状",
                "required_fields": [
                    {"key": "plaintiff_name", "label": "原告姓名/名称", "type": "text", "required": True},
                    {"key": "plaintiff_id", "label": "原告身份证号/统一社会信用代码", "type": "text", "required": True},
                    {"key": "plaintiff_address", "label": "原告住所地", "type": "text", "required": True},
                    {"key": "plaintiff_phone", "label": "原告联系电话", "type": "text", "required": True},
                    {"key": "defendant_name", "label": "被告姓名/名称", "type": "text", "required": True},
                    {"key": "defendant_address", "label": "被告住所地", "type": "text", "required": True},
                    {"key": "contract_type", "label": "合同类型", "type": "text", "required": True},
                    {"key": "contract_date", "label": "合同签订日期", "type": "text", "required": True},
                    {"key": "claim_amount", "label": "诉讼请求金额（元）", "type": "number", "required": True},
                    {"key": "facts", "label": "事实与理由", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事起诉状

原告：{{plaintiff_name}}
身份证号/统一社会信用代码：{{plaintiff_id}}
住所地：{{plaintiff_address}}
联系电话：{{plaintiff_phone}}

被告：{{defendant_name}}
住所地：{{defendant_address}}

诉讼请求：

1. 判令被告支付{{contract_type}}项下款项人民币{{claim_amount}}元；
2. 判令被告支付逾期付款违约金（以{{claim_amount}}元为基数，自应付款之日起按LPR的1.5倍计算至实际清偿之日止）；
3. 判令被告承担本案全部诉讼费用。

事实与理由：

{{facts}}

综上所述，被告的行为已构成违约，严重侵害了原告的合法权益。为此，原告依据《中华人民共和国民法典》第五百七十七条、第五百八十四条及《中华人民共和国民事诉讼法》第一百一十九条之规定，特向贵院提起诉讼，恳请贵院依法审理，支持原告全部诉讼请求。

此致
{{court_name}}

具状人：{{plaintiff_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "civil_complaint_loan",
                "category": "起诉类",
                "title": "民事起诉状（民间借贷）",
                "description": "适用于自然人之间或自然人与企业之间的借款纠纷",
                "required_fields": [
                    {"key": "plaintiff_name", "label": "原告姓名", "type": "text", "required": True},
                    {"key": "plaintiff_id", "label": "原告身份证号", "type": "text", "required": True},
                    {"key": "plaintiff_address", "label": "原告住所地", "type": "text", "required": True},
                    {"key": "plaintiff_phone", "label": "原告联系电话", "type": "text", "required": True},
                    {"key": "defendant_name", "label": "被告姓名", "type": "text", "required": True},
                    {"key": "defendant_address", "label": "被告住所地", "type": "text", "required": True},
                    {"key": "loan_date", "label": "借款日期", "type": "text", "required": True},
                    {"key": "loan_amount", "label": "借款金额（元）", "type": "number", "required": True},
                    {"key": "agreed_interest", "label": "约定利率", "type": "text", "required": True},
                    {"key": "repaid_amount", "label": "已还金额（元）", "type": "number", "required": False},
                    {"key": "facts", "label": "借款经过", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事起诉状

原告：{{plaintiff_name}}
身份证号：{{plaintiff_id}}
住所地：{{plaintiff_address}}
联系电话：{{plaintiff_phone}}

被告：{{defendant_name}}
住所地：{{defendant_address}}

诉讼请求：

1. 判令被告立即偿还原告借款本金人民币{{loan_amount}}元及利息（自{{loan_date}}起按{{agreed_interest}}计算至实际清偿之日止）；
2. 判令被告承担本案全部诉讼费用。

事实与理由：

{{facts}}

综上所述，被告拒不偿还借款的行为严重侵害了原告的合法权益。依据《中华人民共和国民法典》第六百七十五条、第六百七十六条及《最高人民法院关于审理民间借贷案件适用法律若干问题的规定》，特向贵院提起诉讼。

此致
{{court_name}}

具状人：{{plaintiff_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "civil_complaint_labor",
                "category": "起诉类",
                "title": "民事起诉状（劳动争议）",
                "description": "适用于追索劳动报酬、违法解除劳动合同赔偿等劳动争议",
                "required_fields": [
                    {"key": "plaintiff_name", "label": "原告姓名", "type": "text", "required": True},
                    {"key": "plaintiff_id", "label": "原告身份证号", "type": "text", "required": True},
                    {"key": "defendant_name", "label": "被告公司全称", "type": "text", "required": True},
                    {"key": "defendant_address", "label": "被告住所地", "type": "text", "required": True},
                    {"key": "employment_start", "label": "入职日期", "type": "text", "required": True},
                    {"key": "monthly_salary", "label": "月工资（元）", "type": "number", "required": True},
                    {"key": "claim_amount", "label": "诉求金额（元）", "type": "number", "required": True},
                    {"key": "facts", "label": "事实与理由", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事起诉状

原告：{{plaintiff_name}}
身份证号：{{plaintiff_id}}

被告：{{defendant_name}}
住所地：{{defendant_address}}

诉讼请求：

1. 判令被告向原告支付人民币{{claim_amount}}元；
2. 判令被告承担本案全部诉讼费用。

事实与理由：

原告于{{employment_start}}入职被告处，月工资{{monthly_salary}}元。

{{facts}}

综上所述，被告的行为违反了《中华人民共和国劳动合同法》的相关规定，原告为维护自身合法权益，特向贵院提起诉讼。

此致
{{court_name}}

具状人：{{plaintiff_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "civil_complaint_divorce",
                "category": "起诉类",
                "title": "民事起诉状（离婚纠纷）",
                "description": "适用于离婚诉讼，含财产分割、子女抚养等诉求",
                "required_fields": [
                    {"key": "plaintiff_name", "label": "原告姓名", "type": "text", "required": True},
                    {"key": "plaintiff_id", "label": "原告身份证号", "type": "text", "required": True},
                    {"key": "defendant_name", "label": "被告姓名", "type": "text", "required": True},
                    {"key": "defendant_id", "label": "被告身份证号", "type": "text", "required": True},
                    {"key": "marriage_date", "label": "结婚日期", "type": "text", "required": True},
                    {"key": "children_info", "label": "子女情况（如无则填'无'）", "type": "textarea", "required": True},
                    {"key": "property_info", "label": "共同财产情况", "type": "textarea", "required": True},
                    {"key": "facts", "label": "离婚理由", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事起诉状

原告：{{plaintiff_name}}
身份证号：{{plaintiff_id}}

被告：{{defendant_name}}
身份证号：{{defendant_id}}

诉讼请求：

1. 判令原告与被告离婚；
2. 判令婚生子女由原告/被告抚养（以有利于子女成长为原则）；
3. 判令依法分割夫妻共同财产；
4. 判令被告承担本案诉讼费用。

事实与理由：

原告与被告于{{marriage_date}}登记结婚。

子女情况：{{children_info}}

共同财产：{{property_info}}

{{facts}}

综上所述，夫妻感情确已破裂，无和好可能。依据《中华人民共和国民法典》第一千零七十九条之规定，特向贵院提起诉讼，恳请依法判决。

此致
{{court_name}}

具状人：{{plaintiff_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "property_preservation",
                "category": "申请类",
                "title": "财产保全申请书",
                "description": "适用于诉讼前或诉讼中申请查封、冻结对方财产",
                "required_fields": [
                    {"key": "applicant_name", "label": "申请人姓名/名称", "type": "text", "required": True},
                    {"key": "respondent_name", "label": "被申请人姓名/名称", "type": "text", "required": True},
                    {"key": "preservation_amount", "label": "申请保全金额（元）", "type": "number", "required": True},
                    {"key": "preservation_target", "label": "保全标的（如银行账户/房产/车辆等）", "type": "textarea", "required": True},
                    {"key": "reason", "label": "申请理由", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """财产保全申请书

申请人：{{applicant_name}}

被申请人：{{respondent_name}}

申请事项：

请求贵院依法对被申请人名下价值人民币{{preservation_amount}}元的财产采取保全措施，保全标的为：
{{preservation_target}}

事实与理由：

{{reason}}

申请人深知，如不立即采取财产保全措施，被申请人可能转移、隐匿财产，导致将来判决难以执行。为此，申请人依据《中华人民共和国民事诉讼法》第一百零三条、第一百零四条之规定，特向贵院申请财产保全。

申请人愿提供相应担保。

此致
{{court_name}}

申请人：{{applicant_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "enforcement_application",
                "category": "申请类",
                "title": "强制执行申请书",
                "description": "适用于生效判决/调解书/仲裁裁决的强制执行申请",
                "required_fields": [
                    {"key": "applicant_name", "label": "申请人姓名/名称", "type": "text", "required": True},
                    {"key": "respondent_name", "label": "被申请人姓名/名称", "type": "text", "required": True},
                    {"key": "judgment_info", "label": "生效法律文书信息（案号+日期）", "type": "text", "required": True},
                    {"key": "execution_amount", "label": "申请执行金额（元）", "type": "number", "required": True},
                    {"key": "facts", "label": "未履行情况说明", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """强制执行申请书

申请人：{{applicant_name}}

被申请人：{{respondent_name}}

申请事项：

请求贵院依法强制执行{{judgment_info}}，责令被申请人向申请人支付人民币{{execution_amount}}元。

事实与理由：

{{facts}}

被申请人未在法定期限内履行生效法律文书确定的义务，申请人依据《中华人民共和国民事诉讼法》第二百四十三条之规定，特向贵院申请强制执行。

此致
{{court_name}}

申请人：{{applicant_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "labor_arbitration",
                "category": "申请类",
                "title": "劳动仲裁申请书",
                "description": "适用于向劳动仲裁委员会申请劳动仲裁",
                "required_fields": [
                    {"key": "applicant_name", "label": "申请人姓名", "type": "text", "required": True},
                    {"key": "applicant_id", "label": "申请人身份证号", "type": "text", "required": True},
                    {"key": "applicant_phone", "label": "申请人联系电话", "type": "text", "required": True},
                    {"key": "respondent_name", "label": "被申请人公司全称", "type": "text", "required": True},
                    {"key": "respondent_address", "label": "被申请人住所地", "type": "text", "required": True},
                    {"key": "employment_start", "label": "入职日期", "type": "text", "required": True},
                    {"key": "claim_items", "label": "仲裁请求（逐条列明）", "type": "textarea", "required": True},
                    {"key": "facts", "label": "事实与理由", "type": "textarea", "required": True},
                    {"key": "committee_name", "label": "呈送仲裁委员会名称", "type": "text", "required": True},
                ],
                "template_content": """劳动仲裁申请书

申请人：{{applicant_name}}
身份证号：{{applicant_id}}
联系电话：{{applicant_phone}}

被申请人：{{respondent_name}}
住所地：{{respondent_address}}

仲裁请求：

{{claim_items}}

事实与理由：

申请人于{{employment_start}}入职被申请人处。

{{facts}}

综上所述，被申请人的行为违反了《中华人民共和国劳动合同法》的相关规定，申请人依据《中华人民共和国劳动争议调解仲裁法》第二十七条、第二十八条之规定，特向贵委申请仲裁。

此致
{{committee_name}}

申请人：{{applicant_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "civil_answer",
                "category": "答辩类",
                "title": "民事答辩状",
                "description": "适用于被告对原告起诉状的答辩回应",
                "required_fields": [
                    {"key": "respondent_name", "label": "答辩人姓名/名称", "type": "text", "required": True},
                    {"key": "case_info", "label": "针对的案件信息（案号+原告）", "type": "text", "required": True},
                    {"key": "defense_points", "label": "答辩要点（逐条反驳）", "type": "textarea", "required": True},
                    {"key": "conclusion", "label": "答辩结论", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事答辩状

答辩人：{{respondent_name}}

答辩人就{{case_info}}一案，提出答辩如下：

{{defense_points}}

综上所述，{{conclusion}}

此致
{{court_name}}

答辩人：{{respondent_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "civil_appeal",
                "category": "上诉类",
                "title": "民事上诉状",
                "description": "适用于不服一审判决/裁定的上诉",
                "required_fields": [
                    {"key": "appellant_name", "label": "上诉人姓名/名称", "type": "text", "required": True},
                    {"key": "original_judgment", "label": "原审判决信息（法院+案号+日期）", "type": "text", "required": True},
                    {"key": "appeal_requests", "label": "上诉请求", "type": "textarea", "required": True},
                    {"key": "appeal_reasons", "label": "上诉理由", "type": "textarea", "required": True},
                    {"key": "court_name", "label": "呈送上级法院名称", "type": "text", "required": True},
                ],
                "template_content": """民事上诉状

上诉人：{{appellant_name}}

上诉人因不服{{original_judgment}}，现提出上诉。

上诉请求：

{{appeal_requests}}

上诉理由：

{{appeal_reasons}}

此致
{{court_name}}

上诉人：{{appellant_name}}
    年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "divorce_agreement",
                "category": "协议类",
                "title": "离婚协议书",
                "description": "适用于夫妻双方自愿协议离婚",
                "required_fields": [
                    {"key": "party_a", "label": "男方姓名", "type": "text", "required": True},
                    {"key": "party_a_id", "label": "男方身份证号", "type": "text", "required": True},
                    {"key": "party_b", "label": "女方姓名", "type": "text", "required": True},
                    {"key": "party_b_id", "label": "女方身份证号", "type": "text", "required": True},
                    {"key": "marriage_date", "label": "结婚日期", "type": "text", "required": True},
                    {"key": "children_arrangement", "label": "子女抚养安排", "type": "textarea", "required": True},
                    {"key": "property_arrangement", "label": "财产分割安排", "type": "textarea", "required": True},
                    {"key": "debt_arrangement", "label": "债务处理安排", "type": "textarea", "required": True},
                ],
                "template_content": """离婚协议书

男方：{{party_a}}
身份证号：{{party_a_id}}

女方：{{party_b}}
身份证号：{{party_b_id}}

男女双方于{{marriage_date}}登记结婚。现双方因感情破裂，自愿协议离婚，经协商一致，达成如下协议：

一、自愿离婚
双方自愿解除婚姻关系。

二、子女抚养
{{children_arrangement}}

三、财产分割
{{property_arrangement}}

四、债务处理
{{debt_arrangement}}

五、其他
本协议一式三份，男女双方各执一份，婚姻登记机关存档一份。
本协议自双方签字并在婚姻登记机关办理离婚登记之日起生效。

男方签字：            女方签字：

    年    月    日          年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "settlement_agreement",
                "category": "协议类",
                "title": "和解协议书",
                "description": "适用于纠纷双方协商达成和解",
                "required_fields": [
                    {"key": "party_a", "label": "甲方姓名/名称", "type": "text", "required": True},
                    {"key": "party_b", "label": "乙方姓名/名称", "type": "text", "required": True},
                    {"key": "dispute_description", "label": "争议事项简述", "type": "textarea", "required": True},
                    {"key": "settlement_terms", "label": "和解条款", "type": "textarea", "required": True},
                    {"key": "settlement_date", "label": "签署日期", "type": "text", "required": True},
                ],
                "template_content": """和解协议书

甲方：{{party_a}}
乙方：{{party_b}}

鉴于：
{{dispute_description}}

现甲乙双方经友好协商，自愿达成如下和解协议：

{{settlement_terms}}

本协议一式两份，甲乙双方各执一份，自双方签字（盖章）之日起生效。
本协议的签订、履行及解释均适用中华人民共和国法律。
因本协议引起的任何争议，双方应协商解决；协商不成的，任何一方有权向甲方所在地有管辖权的人民法院提起诉讼。

甲方（签字/盖章）：      乙方（签字/盖章）：

    年    月    日            年    月    日""",
                "download_count": 0,
            },
            {
                "template_id": "demand_letter",
                "category": "函件类",
                "title": "律师函（催告）",
                "description": "适用于委托律师向对方发送正式催告函",
                "required_fields": [
                    {"key": "sender_name", "label": "委托人姓名/名称", "type": "text", "required": True},
                    {"key": "recipient_name", "label": "收件人姓名/名称", "type": "text", "required": True},
                    {"key": "law_firm", "label": "律师事务所名称", "type": "text", "required": True},
                    {"key": "matter_description", "label": "委托事项描述", "type": "textarea", "required": True},
                    {"key": "demands", "label": "具体诉求", "type": "textarea", "required": True},
                    {"key": "deadline", "label": "履行期限", "type": "text", "required": True},
                    {"key": "consequences", "label": "逾期后果", "type": "textarea", "required": True},
                ],
                "template_content": """律 师 函

（ ）{{law_firm}}律函字第  号

致：{{recipient_name}}

{{law_firm}}接受{{sender_name}}的委托，就{{matter_description}}一事，特向贵方致函如下：

{{demands}}

请贵方于{{deadline}}前履行上述义务。如逾期未履行，本所将根据委托人的授权，依法采取包括但不限于提起诉讼、申请财产保全等法律措施，由此产生的一切法律后果由贵方自行承担。

{{consequences}}

特此函告。

{{law_firm}}
律师：
    年    月    日""",
                "download_count": 0,
            },
        ]
        self._store_path.write_text(json.dumps(templates, ensure_ascii=False, indent=2), encoding="utf-8")

    def _load(self) -> list:
        if not self._store_path.exists():
            return []
        return json.loads(self._store_path.read_text(encoding="utf-8"))

    def _save(self, data: list):
        self._store_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def list_templates(self) -> list[dict]:
        data = self._load()
        return [{
            "template_id": t["template_id"],
            "category": t["category"],
            "title": t["title"],
            "description": t.get("description", ""),
            "required_fields": t.get("required_fields", []),
            "download_count": t.get("download_count", 0),
        } for t in data]

    def get_template(self, template_id: str) -> dict | None:
        data = self._load()
        for t in data:
            if t["template_id"] == template_id:
                return t
        return None

    def assemble(self, template_id: str, fields: dict) -> dict | None:
        """用字段值填充模板，生成文书并导出为 docx"""
        t = self.get_template(template_id)
        if not t:
            return None

        content = t["template_content"]
        for key, val in fields.items():
            content = content.replace("{{" + key + "}}", str(val or ""))

        # 更新下载计数
        data = self._load()
        for d in data:
            if d["template_id"] == template_id:
                d["download_count"] = d.get("download_count", 0) + 1
                break
        self._save(data)

        # 生成 docx
        filename = f"{template_id}_{uuid.uuid4().hex[:6]}.docx"
        filepath = self._export_dir / filename
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)

        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5

        doc.save(str(filepath))

        # 未填的占位符保留提示
        for key in t.get("required_fields", []):
            content = content.replace("{{" + key["key"] + "}}", f"[未填写：{key['label']}]")

        return {
            "content": content,
            "download_url": f"/api/templates/{template_id}/download/{filename}",
            "filename": filename,
        }

    def get_template_download(self, template_id: str) -> Path | None:
        """返回空白模板下载路径"""
        t = self.get_template(template_id)
        if not t:
            return None

        filename = f"{template_id}_template.docx"
        filepath = self._export_dir / filename

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)

        content = t["template_content"]
        for f in t.get("required_fields", []):
            content = content.replace("{{" + f["key"] + "}}", f"【{f['label']}】")

        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5

        doc.save(str(filepath))
        return filepath


document_service = DocumentService()
