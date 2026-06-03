"""
法律文档解析器 —— 支持 PDF/DOCX/TXT，条款级切分
"""
import re
from pathlib import Path
from typing import Iterator


class LegalDocumentParser:
    """法律文档解析器，按条款级切分"""

    @staticmethod
    def parse(file_path: str) -> list[dict]:
        """解析文档，返回条款列表"""
        ext = Path(file_path).suffix.lower()
        if ext == ".txt" or ext == ".md":
            return LegalDocumentParser._parse_text(file_path)
        elif ext == ".docx":
            return LegalDocumentParser._parse_docx(file_path)
        elif ext == ".pdf":
            return LegalDocumentParser._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _parse_text(file_path: str) -> list[dict]:
        """解析 TXT/MD 格式的法律文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        chunks = []
        lines = content.split("\n")
        current_law = ""
        current_chapter = ""
        current_article = ""
        article_content: list[str] = []
        metadata = {"effective_date": "", "status": ""}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 法律名称
            if line.startswith("# ") and not line.startswith("## "):
                current_law = line[2:].strip()
                continue

            # 元数据
            if line.startswith("## 施行日期:"):
                metadata["effective_date"] = line.replace("## 施行日期:", "").strip()
                continue
            if line.startswith("## 状态:"):
                metadata["status"] = line.replace("## 状态:", "").strip()
                continue

            # 章节
            if line.startswith("## ") and "第" not in line:
                # 结束前一条款
                if current_article and article_content:
                    chunks.append(
                        {
                            "law": current_law,
                            "chapter": current_chapter,
                            "article": current_article,
                            "content": "\n".join(article_content),
                            "effective_date": metadata.get("effective_date", ""),
                            "status": metadata.get("status", ""),
                        }
                    )
                    article_content = []
                    current_article = ""
                current_chapter = line[3:].strip()
                continue

            # 条款（支持 markdown ###/## 标题 和 普通"第X条"格式）
            is_article_header = (
                line.startswith("### ")
                or (line.startswith("## ") and "第" in line[:10])
                or bool(re.match(r"第[一二三四五六七八九十百千\d]+条", line[:10]))
            )
            if is_article_header:
                if current_article and article_content:
                    chunks.append(
                        {
                            "law": current_law,
                            "chapter": current_chapter,
                            "article": current_article,
                            "content": "\n".join(article_content),
                            "effective_date": metadata.get("effective_date", ""),
                            "status": metadata.get("status", ""),
                        }
                    )
                    article_content = []
                clean = line.replace("### ", "").replace("## ", "").strip()
                current_article = clean
                continue

            article_content.append(line)

        # 最后一条
        if current_article and article_content:
            chunks.append(
                {
                    "law": current_law,
                    "chapter": current_chapter,
                    "article": current_article,
                    "content": "\n".join(article_content),
                    "effective_date": metadata.get("effective_date", ""),
                    "status": metadata.get("status", ""),
                }
            )

        # 兜底：没识别到任何条款时，全文作为一个块
        if not chunks:
            chunks = [{"article": "全文", "content": content, "law": current_law or Path(file_path).stem}]

        return chunks

    @staticmethod
    def _parse_docx(file_path: str) -> list[dict]:
        """解析 DOCX 合同文件"""
        from docx import Document

        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        # 按段落切分，尝试识别条款号
        chunks = []
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        current_clause = ""
        clause_content: list[str] = []

        for para in paragraphs:
            # 识别条款号（如 第一条、第1条、1. 等）
            if para.startswith("第") or (para[0].isdigit() and ("." in para[:5] or "、" in para[:5])):
                if current_clause:
                    chunks.append({"title": current_clause, "content": "\n".join(clause_content)})
                    clause_content = []
                current_clause = para
            else:
                clause_content.append(para)

        if current_clause:
            chunks.append({"title": current_clause, "content": "\n".join(clause_content)})

        if not chunks:
            # 无法识别条款结构，整篇作为一个块
            chunks = [{"title": Path(file_path).stem, "content": full_text}]

        return chunks

    @staticmethod
    def _parse_pdf(file_path: str) -> list[dict]:
        """解析 PDF 文件"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            full_text = "\n".join(text_parts)
        except Exception:
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
                full_text = "\n".join(text_parts)
            except Exception as e:
                raise ValueError(f"PDF 解析失败: {e}")

        # 按条款切分
        lines = full_text.split("\n")
        chunks = []
        current_title = ""
        current_lines: list[str] = []

        for line in lines:
            line = line.strip()
            if not line:
                if current_lines:
                    current_lines.append("")
                continue
            if line.startswith("第") and ("条" in line[:20]):
                if current_title or current_lines:
                    chunks.append(
                        {"title": current_title or Path(file_path).stem, "content": "\n".join(current_lines)}
                    )
                    current_lines = []
                current_title = line
            else:
                current_lines.append(line)

        if current_title or current_lines:
            chunks.append(
                {"title": current_title or Path(file_path).stem, "content": "\n".join(current_lines)}
            )

        if not chunks:
            chunks = [{"title": Path(file_path).stem, "content": full_text}]

        return chunks

    @classmethod
    def parse_to_documents(cls, file_path: str) -> list[dict]:
        """解析文档并添加元数据，用于 RAG 索引"""
        chunks = cls.parse(file_path)
        source = Path(file_path).stem
        for i, chunk in enumerate(chunks):
            chunk["source"] = source
            chunk["chunk_id"] = f"{source}_{i}"
        return chunks


legal_parser = LegalDocumentParser()
